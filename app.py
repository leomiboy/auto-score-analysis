import streamlit as st
import pandas as pd
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
import time

# ==========================================
# 專案：班級讀書建議生成器 (最終完整版 v4)
# 功能：
# 1. 讀取 Excel (5分頁)
# 2. 自動篩選 API Key 可用的 Text-out 模型
# 3. 新增 gemini-3-flash-preview 至推薦清單
# 4. AI 生成建議 (使用新版 GEM 嚴格提示詞)
# 5. 產出 Word 檔 (.docx)
# ==========================================

# --- 1. 網頁設定 ---
st.set_page_config(page_title="班級讀書建議生成器", layout="wide")
st.title("🎓 班級錯題分析與讀書建議生成器 (Word版)")
st.markdown("""
此工具協助老師快速生成全班學生的個別化讀書建議 **Word 檔**。
1. 輸入您的 **Gemini API Key** (系統會自動搜尋可用的文字生成模型)。
2. 選擇合適的 AI 模型 (請參考下方說明)。
3. 系統將自動分析並打包 ZIP 下載。
""")

# --- 新增：模型說明表格 ---
st.markdown("### 🤖 模型選擇指南")
st.markdown("""
| 模型名稱 | 適用場景 | 模型特點 | 成績分析推薦 |
| :--- | :--- | :--- | :--- |
| **Gemini 3 Flash (Preview)** | 複雜邏輯、代碼撰寫、多步驟規劃、學術論文分析 | **最強大**。具備「Deep Think」能力，適合需要深度推理的高難度任務。 | **深度學情診斷**：適合分析全校或跨年級的長期成績趨勢，找出隱藏的學習斷點，並生成個人化的深度學習改進策略建議。 |
| **Gemini 2.5 Flash** | 日常問答、一般文章摘要、圖片理解、數據提取 | **主力機**。效能最均衡，是處理多數中階任務的首選。 | **標準化成績報告**：適合將班級考卷分數轉化為視覺化摘要，撰寫每位學生的進步評語，或從手寫成績單照片中提取數據。 |
| **Gemini 2.5 Flash-Lite** | 簡單分類、翻譯、大量垃圾郵件過濾、標籤生成 | **極速且省錢**：延遲最低，最適合處理大量簡單、重複性的自動化任務。 | **基礎數據整理**：適合執行大規模的分數排序、及格/不及格分類、或是將原始成績轉換為簡單的等第（如 A、B、C）。 |
""")
st.markdown("---")

# --- 2. 核心邏輯函式 ---

def get_available_models(api_key):
    """
    使用使用者的 API Key 查詢 Google 帳號下可用的模型列表
    並進行嚴格篩選與排序：推薦模型置頂
    """
    try:
        genai.configure(api_key=api_key)
        
        # 定義推薦模型清單 (已加入 gemini-3-flash-preview)
        recommended_models = [
            "gemini-3-flash-preview",
            "gemini-3-flash", 
            "gemini-2.5-flash", 
            "gemini-2.5-flash-lite"
        ]
        
        all_models = []
        for m in genai.list_models():
            # 條件 1: 必須支援 'generateContent'
            if 'generateContent' in m.supported_generation_methods:
                # 條件 2: 必須是 Gemini 系列
                if 'gemini' in m.name:
                    # 條件 3: 排除 Vision/Embedding
                    if 'vision' not in m.name and 'embedding' not in m.name:
                        clean_name = m.name.replace('models/', '')
                        all_models.append(clean_name)
        
        final_list = []
        
        # 先加入推薦模型 (如果使用者帳號有權限的話)
        for rec in recommended_models:
            if rec in all_models:
                final_list.append(rec)
                all_models.remove(rec) # 避免重複
        
        # 加入剩下的模型 (包含 1.5, 2.0 等，按名稱排序)
        all_models.sort(reverse=True)
        final_list.extend(all_models)
        
        return final_list
    except Exception as e:
        return []

def format_model_name(model_name):
    """格式化模型名稱顯示"""
    # 定義哪些模型要顯示推薦標籤
    recommended_set = [
        "gemini-3-flash-preview",
        "gemini-3-flash", 
        "gemini-2.5-flash", 
        "gemini-2.5-flash-lite"
    ]
    
    if model_name in recommended_set:
        return f"{model_name} (推薦 🔥)"
    return model_name

def process_excel_data(uploaded_file):
    """讀取 Excel 並整理所有學生的錯題"""
    try:
        xls = pd.ExcelFile(uploaded_file)
    except Exception:
        return None, "檔案格式錯誤，請確認上傳的是 .xlsx Excel 檔案。"

    required_sheets = ["國文", "英文", "數學", "社會", "自然"]
    missing_sheets = [sheet for sheet in required_sheets if sheet not in xls.sheet_names]
    
    if missing_sheets:
        return None, f"Excel 缺少必要分頁：{missing_sheets}，請確認分頁名稱正確。"

    data_map = {}
    for sheet in required_sheets:
        data_map[sheet] = pd.read_excel(xls, sheet_name=sheet, header=None)

    try:
        first_df = data_map["國文"]
        student_list = first_df.iloc[5:, 1].dropna().unique().tolist()
    except Exception as e:
        return None, f"無法讀取學生名單，請確認 Excel 格式 (錯誤訊息: {e})"
    
    all_students_data = {}
    
    for student in student_list:
        student_errors = {}
        for subject in required_sheets:
            df = data_map[subject]
            try:
                q_nums = df.iloc[0, 2:].values
                categories = df.iloc[1, 2:].values
                k_points = df.iloc[2, 2:].values
                
                student_df_temp = df.iloc[5:, 1:].reset_index(drop=True)
                student_df_temp.columns = ["Name"] + [i for i in range(len(student_df_temp.columns)-1)]
                
                target_row = student_df_temp[student_df_temp["Name"] == student]
                
                if target_row.empty:
                    continue
                
                answers = target_row.iloc[0, 1:].values
                
                errors = []
                for ans, cat, kp, qn in zip(answers, categories, k_points, q_nums):
                    ans_str = str(ans).strip()
                    if ans_str != "-" and pd.notna(ans) and ans_str != "":
                        errors.append({
                            "題號": qn,
                            "領域": str(cat).strip() if pd.notna(cat) else "其他",
                            "知識點": kp
                        })
                student_errors[subject] = errors
            except Exception as e:
                print(f"處理 {student} 的 {subject} 時發生錯誤: {e}")
                
        all_students_data[student] = student_errors
        
    return all_students_data, None

def get_ai_advice(api_key, model_name, student_name, error_data):
    """呼叫 Gemini 生成建議 (使用新版 GEM 嚴格提示詞)"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        # ==========================================
        # ⚠️ 新版 GEM 嚴格提示詞 (一字不漏置入)
        # ==========================================
        prompt = f"""
        學生姓名：{student_name}
        錯題數據：{error_data}

        你是一位專業的台灣國中教育會考升學輔導專家。你的任務是讀取使用者上傳的 JSON 格式錯題分析檔案（九年級第2次複習考，範圍1-4冊），並生成一份精準的讀書建議報告。

        請嚴格遵守以下規則進行分析與輸出：

        ### 核心規則 (Critical Rules)
        1.  **直接開始**：**絕對不要**有任何開場白（如「親愛的同學你好」、「我是你的分析師」）。請直接以「## 一、 【整體表現總評】」作為輸出的第一行。
        2.  **統一稱呼**：報告中若需提及學生，請一律使用代名詞**「你」**，**嚴禁**使用 JSON 檔案中的學生姓名。
        3.  **無結尾提問**：報告結束時，請給予一句簡短的鼓勵即可，**不要**詢問「是否需要更多協助」。
        4.  **格式一致性**：必須嚴格依照下方的【輸出範本】格式進行排版。

        ### 步驟一：資料分類 (Knowledge Segmentation)
        請根據 JSON 中的「知識點」敘述，運用學科專業將錯題歸類到以下領域：

        *   **國文科**：
            *   【文言文】：古文、詩詞曲、文言閱讀、形音義(古字)、國學常識。
            *   【白話文】：現代散文、小說、新詩、白話閱讀、應用文、標點。
        *   **英文科**：
            *   【聽力】：聽力測驗、辨識句意、言談理解。
            *   【閱讀】：單字、文法、閱讀測驗、克漏字。
        *   **數學科**：
            *   【代數】：數與量、多項式、方程式、函數、不等式、數列級數。
            *   【幾何】：圖形性質、三角形(全等/相似)、四邊形、圓形、尺規作圖、畢氏定理。
            *   【機率統計】：資料分析、圖表判讀、機率。
        *   **社會科**：
            *   【歷史】：台灣史、中國史、世界史。
            *   【地理】：地理概說、區域地理(台/中/世)、氣候、地形、水文、產業。
            *   【公民】：政治、法律、經濟、社會文化。
        *   **自然科**：
            *   【生物】：細胞、生理、遺傳、演化、生態 (七年級範圍)。
            *   【理化】：物理(力/電/熱/光/波)、化學(反應/元素/酸鹼) (八九年級範圍)。
            *   【地科】：天文、地質、氣象、海洋 (九年級範圍)。

        ### 步驟二：分析邏輯
        1.  **強弱科判斷**：錯題數最少的 1-2 科為「穩定發展科（強科）」；錯題數最多的 1-2 科為「急需搶救科（弱科）」。
        2.  **弱點診斷**：每科找出錯題數最多的前 3 個知識點。
        3.  **領域佔比計算**：計算該科錯題在上述分類的百分比（例如：該科錯10題，代數錯6題，則代數佔60%）。

        ### 步驟三：輸出範本 (Output Template)
        請完全依照以下 Markdown 結構輸出內容：

        ## 一、 【整體表現總評】

        * **強弱科分析**：
            * **穩定發展科（強科）**：**[科目名]**（[錯題數]題）。[簡短評語]
            * **急需搶救科（弱科）**：**[科目名]**（[錯題數]題）。[簡短評語]

        * **關鍵弱點領域**：
        [跨科目分析該生的痛點。例如：你的痛點主要在於「高難度的邏輯推演」。相較之下，純記憶性的考題對你而言困難度較低，但如何將點狀的知識連成線，是接下來衝刺期的關鍵。]

        ---

        ## 二、 【分科深度分析與建議】

        ### 1. 國文科：[請給予一句該科的總結短評]

        * **弱點診斷 (前三名)**：
            1. **【[領域名]】** [知識點名稱]
            2. **【[領域名]】** [知識點名稱]
            3. **【[領域名]】** [知識點名稱]

        * **領域佔比分析**：
            *   **文言文**：[X]%
            *   **白話文**：[Y]%
            *(請確保加總約為 100%)*

        * **會考衝刺建議**：
        [針對弱點提供具體讀書建議。]

        ### 2. 英文科：[請給予一句該科的總結短評]

        * **弱點診斷 (前三名)**：
            1. **【[領域名]】** [知識點名稱]
            2. **【[領域名]】** [知識點名稱]
            3. **【[領域名]】** [知識點名稱]

        * **領域佔比分析**：
            *   **聽力**：[X]%
            *   **閱讀**：[Y]%

        * **會考衝刺建議**：
        [針對弱點提供具體讀書建議。]

        ### 3. 數學科：[請給予一句該科的總結短評]

        * **弱點診斷 (前三名)**：
            1. **【[領域名]】** [知識點名稱]
            2. **【[領域名]】** [知識點名稱]
            3. **【[領域名]】** [知識點名稱]

        * **領域佔比分析**：
            *   **代數**：[X]%
            *   **幾何**：[Y]%
            *   **機率統計**：[Z]%

        * **會考衝刺建議**：
        [針對弱點提供具體讀書建議。]

        ### 4. 社會科：[請給予一句該科的總結短評]

        * **弱點診斷 (前三名)**：
            1. **【[領域名]】** [知識點名稱]
            2. **【[領域名]】** [知識點名稱]
            3. **【[領域名]】** [知識點名稱]

        * **領域佔比分析**：
            *   **歷史**：[X]%
            *   **地理**：[Y]%
            *   **公民**：[Z]%

        * **會考衝刺建議**：
        [針對弱點提供具體讀書建議。]

        ### 5. 自然科：[請給予一句該科的總結短評]

        * **弱點診斷 (前三名)**：
            1. **【[領域名]】** [知識點名稱]
            2. **【[領域名]】** [知識點名稱]
            3. **【[領域名]】** [知識點名稱]

        * **領域佔比分析**：
            *   **生物**：[X]%
            *   **理化**：[Y]%
            *   **地科**：[Z]%

        * **會考衝刺建議**：
        [針對弱點提供具體讀書建議。]

        ---

        **[請在此處給予一段總結性的鼓勵話語，強調複習考是為了找洞而非打擊信心，並建議從哪個科目開始著手最有效率。]**
        """
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"AI 分析連線失敗: {e} (請檢查 API Key 或模型權限)"

def create_word(student_name, ai_advice):
    """建立 Word 文件 (.docx)"""
    doc = Document()
    
    title = doc.add_heading(f"{student_name} - 讀書建議報告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    clean_text = ai_advice.replace('**', '').replace('## ', '').replace('### ', '')
    
    for paragraph_text in clean_text.split('\n'):
        if paragraph_text.strip():
            p = doc.add_paragraph(paragraph_text)
            p.style.font.size = Pt(12)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. 介面互動邏輯 ---

with st.sidebar:
    st.header("🔑 設定")
    user_api_key = st.text_input("請輸入 Gemini API Key", type="password", help="請前往 Google AI Studio 申請")
    
    selected_model = None
    
    if user_api_key:
        with st.spinner("正在驗證 Key 並搜尋可用模型..."):
            available_models = get_available_models(user_api_key)
            
        if available_models:
            st.success(f"驗證成功！找到 {len(available_models)} 個可用模型")
            selected_model = st.selectbox(
                "🤖 請選擇 AI 模型", 
                available_models,
                index=0,
                format_func=format_model_name,
                help="已自動過濾掉不支援文字生成的模型，並將推薦模型置頂。"
            )
        else:
            st.error("無法獲取模型列表，請檢查 API Key 是否正確。")
    
    st.markdown("---")
    st.info("💡 提示：請上傳包含 5 個分頁 (國文, 英文, 數學, 社會, 自然) 的 Excel 檔案。")

uploaded_file = st.file_uploader("📂 上傳 Excel 檔案 (.xlsx)", type=['xlsx'])

if uploaded_file and user_api_key and selected_model:
    if st.button("🚀 開始生成全班報告 (Word)"):
        
        status_text = st.empty()
        progress_bar = st.progress(0)
        
        status_text.text("正在讀取 Excel 資料...")
        all_data, error_msg = process_excel_data(uploaded_file)
        
        if error_msg:
            st.error(error_msg)
        else:
            zip_buffer = io.BytesIO()
            total_students = len(all_data)
            
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                for i, (student, errors) in enumerate(all_data.items()):
                    progress = (i + 1) / total_students
                    progress_bar.progress(progress)
                    status_text.text(f"正在分析：{student} ({i+1}/{total_students})...")
                    
                    # 核心修改：加入 Try-Except 防止單一學生失敗導致全崩
                    try:
                        advice = get_ai_advice(user_api_key, selected_model, student, str(errors))
                        word_data = create_word(student, advice)
                        zf.writestr(f"{student}_讀書建議.docx", word_data.getvalue())
                    except Exception as e:
                        st.warning(f"跳過學生 {student}：發生錯誤 ({e})")
                    
                    # 核心修改：延長休息時間至 4 秒 (符合 Free Tier 15 RPM 限制)
                    time.sleep(4)
            
            progress_bar.progress(100)
            status_text.success("✅ 生成完成！")
            
            st.download_button(
                label="📥 下載全班報告 (ZIP)",
                data=zip_buffer.getvalue(),
                file_name="全班讀書建議報告_Word.zip",
                mime="application/zip"
            )

elif uploaded_file and not user_api_key:
    st.warning("請在左側輸入 API Key 才能開始執行。")